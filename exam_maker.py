import pandas as pd
import numpy as np
import datetime
import os
from docx import Document
from functions import *

balance_over_chps = False    # do we want to balance over chapters?
n_questions = 50  # number of questions to sample
split_q = 21  # Question used to split V1 and V2
max_asked = 2  # remove questions that have already been asked more than max_asked times

ch_to_include = [1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16]
UIDs_to_skip = []
UIDs_to_use = []  # Use these exact UIDs

RIR_cols = ['RIR_1617_1', 'RIR_1718_1', 'RIR_1718_2', 'RIR_1819_1', 'RIR_1819_2', 'RIR_1920_1']
q_file = 'Tentamenvragen.xlsx'

if not os.path.isdir("output"):
    os.mkdir("output")

datafile = pd.read_excel(q_file, index_col=None)
datafile = datafile.sort_values(['CHP'])
datafile = datafile[datafile['CHP'].isin(ch_to_include)]
datafile = datafile[~datafile['Q_UID'].isin(UIDs_to_skip)]

# how many times has this question been asked?
datafile['num_been_asked'] = np.nansum(np.isfinite(datafile[RIR_cols]), axis=1)
print(str(max(datafile['num_been_asked'])) + " times asked: " + str(np.sum(datafile['num_been_asked']==max(datafile['num_been_asked']))))
print(str(max(datafile['num_been_asked'])-1) + " times asked: " + str(np.sum(datafile['num_been_asked']==max(datafile['num_been_asked'])-1)))
print(str(max(datafile['num_been_asked'])-2) + " times asked: " + str(np.sum(datafile['num_been_asked']==max(datafile['num_been_asked'])-2)))
print(str(max(datafile['num_been_asked'])-3) + " times asked: " + str(np.sum(datafile['num_been_asked']==max(datafile['num_been_asked'])-3)))
print(str(max(datafile['num_been_asked'])-4) + " times asked: " + str(np.sum(datafile['num_been_asked']==max(datafile['num_been_asked'])-4)))


if len(UIDs_to_use) == 0:
    datafile = datafile[datafile['num_been_asked'] < max_asked]
    uniqueQ = np.unique(datafile['Q_ID'])
    uniqueQ = uniqueQ[~np.isnan(uniqueQ)]

    if balance_over_chps:
        # determine which chapter to draw a question from, makes sure the exam is balanced across chapters
        chapterQ = []
        for i in range(n_questions):
            chapterQ.append(int(np.random.choice(ch_to_include, 1, replace=True)))
        chapterQ = [1,1,1,2,2,2,3,3,3,4,4,4,5,5,5,6,6,6,7,7,7,8,8,8,8,9,9,9,9,10,10,10,11,11,11,12,12,12,13,13,13,14,14,14,15,15,15,16,16,16]
        if len(chapterQ) != n_questions: raise Exception("Number of chapter selections should match number of questions")
        print("Chapter distribution")
        print(chapterQ)
        questions = []
        for chp in np.unique(chapterQ):
            qs_from_chp = chapterQ.count(chp)
            chp_subset = datafile[datafile['CHP'] == chp]
            uniqueQ = np.unique(chp_subset['Q_ID'])
            uniqueQ = uniqueQ[~np.isnan(uniqueQ)]
            print("Drawing " + str(qs_from_chp) + " questions from chapter " + str(chp) + " where IDs are " + str(uniqueQ))
            questions.extend( np.random.choice(uniqueQ, qs_from_chp, replace=False) )
        # questions = questions[~np.isnan(questions)]
        # orderedQ = sorted(questions, key=lambda x: uniqueQ.tolist().index(x))
        orderedQ = questions
    else:  # we just sample randomly over all Q_IDs
        questions = np.random.choice(uniqueQ, n_questions, replace=False)
        questions = questions[~np.isnan(questions)]
        orderedQ = sorted(questions, key=lambda x: uniqueQ.tolist().index(x))

    UIDs_to_use = []
    for i in orderedQ:
        IDq = datafile[datafile['Q_ID'] == i].sample(1).reset_index(drop=True)
        # IDq = IDq.reset_index(drop=True)
        ID = IDq.loc[0,'Q_UID']
        UIDs_to_use.append(ID)


ord_datafile = datafile[datafile['Q_UID'].isin(UIDs_to_use)]
ord_datafile = ord_datafile.sort_values(['CHP'])
UIDs_to_use = ord_datafile['Q_UID'].tolist()
print("UIDs", UIDs_to_use)
UIDs_v2 = shift(UIDs_to_use, split_q)
print(UIDs_v2)

with open('output/log.txt', 'w') as file_handler:
    d = datetime.datetime.now()
    file_handler.write('Exam generated on {date:%Y-%m-%d} at {date:%H:%M:%S}\n'.format( date=datetime.datetime.now() ))
    file_handler.write("UIDs included:\n")
    file_handler.write(str(UIDs_to_use))


answer_doc = ([['questionV1', 'questionV2', 'answer']])
ans_varsNL = ['A1_NL', 'A2_NL', 'A3_NL', 'A4_NL']
ans_varsEN = ['A1_EN', 'A2_EN', 'A3_EN', 'A4_EN']
ans_let_num = {1: 'A', 2: 'B', 3: 'C', 4: 'D'}
ans_order_dict = {}


# Print questions and their chapters
for i in UIDs_to_use:
    question_data = datafile[datafile['Q_UID']==i].reset_index(drop=True)
    print("Q", i, "  CH", question_data.loc[0,'CHP'])


docNL = Document()
docEN = Document()

q_num = 0
for i in UIDs_to_use:
    q_num += 1
    question_data = datafile[datafile['Q_UID']==i].reset_index(drop=True)

    p = prepare_paragraph(docNL)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(str(q_num) + '.\t')
    write_markdown_paragraph(p, question_data.loc[0,'Q_NL'])

    p = prepare_paragraph(docEN)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(str(q_num) + '.\t')
    write_markdown_paragraph(p, question_data.loc[0,'Q_EN'])

    ans_order = [0, 1, 2, 3]
    if question_data.loc[0,'SHUFFLE_ANSWERS'] == 1:
        ans_order = np.random.choice(ans_order, 4, replace=False)
        print("Shuffling answers for question " + str(i))

    ans_order_dict[i] = ans_order
    print(question_data.loc[0,'COR'])
    cor_resp = int(np.argwhere(ans_order == (question_data.loc[0,'COR'] - 1))) + 1
    cor_resp = ans_let_num[cor_resp]
    print(cor_resp)
    print("q_num", q_num)
    answer_doc.append([q_num, shift(list(range(1, n_questions+1)), -split_q)[q_num-1], cor_resp])


    p = prepare_paragraph(docNL)
    p.add_run(str('A.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[0]]][0])
    p = prepare_paragraph(docEN)
    p.add_run(str('A.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[0]]][0])

    p = prepare_paragraph(docNL)
    p.add_run(str('B.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[1]]][0])
    p = prepare_paragraph(docEN)
    p.add_run(str('B.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[1]]][0])

    p = prepare_paragraph(docNL)
    p.add_run(str('C.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[2]]][0])
    p = prepare_paragraph(docEN)
    p.add_run(str('C.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[2]]][0])

    p = prepare_paragraph(docNL)
    p.paragraph_format.space_after = Pt(24)
    p.add_run(str('D.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[3]]][0])
    p = prepare_paragraph(docEN)
    p.paragraph_format.space_after = Pt(24)
    p.add_run(str('D.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[3]]][0])

docNL.save('output/tentamen_NL_v1.docx')
docEN.save('output/tentamen_EN_v1.docx')


docNL = Document()
docEN = Document()

q_num = 0
for i in UIDs_v2:
    q_num += 1
    question_data = datafile[datafile['Q_UID']==i].reset_index(drop=True)

    p = prepare_paragraph(docNL)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(str(q_num) + '.\t')
    write_markdown_paragraph(p, question_data.loc[0,'Q_NL'])

    p = prepare_paragraph(docEN)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(str(q_num) + '.\t')
    write_markdown_paragraph(p, question_data.loc[0,'Q_EN'])

    ans_order = [0, 1, 2, 3]
    if question_data.loc[0,'SHUFFLE_ANSWERS'] == 1:
        ans_order = ans_order_dict[i]

    p = prepare_paragraph(docNL)
    p.add_run(str('A.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[0]]][0])
    p = prepare_paragraph(docEN)
    p.add_run(str('A.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[0]]][0])

    p = prepare_paragraph(docNL)
    p.add_run(str('B.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[1]]][0])
    p = prepare_paragraph(docEN)
    p.add_run(str('B.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[1]]][0])

    p = prepare_paragraph(docNL)
    p.add_run(str('C.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[2]]][0])
    p = prepare_paragraph(docEN)
    p.add_run(str('C.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[2]]][0])

    p = prepare_paragraph(docNL)
    p.paragraph_format.space_after = Pt(24)
    p.add_run(str('D.\t'))
    write_markdown_paragraph(p, question_data[ans_varsNL[ans_order[3]]][0])
    p = prepare_paragraph(docEN)
    p.paragraph_format.space_after = Pt(24)
    p.add_run(str('D.\t'))
    write_markdown_paragraph(p, question_data[ans_varsEN[ans_order[3]]][0])

docNL.save('output/tentamen_NL_v2.docx')
docEN.save('output/tentamen_EN_v2.docx')

df = pd.DataFrame(answer_doc)
df.to_excel('output/answer_sheet.xlsx', index=False, header=False)