from docx.shared import Inches, Cm, Pt
import xlsxwriter


def shift(l, n):
    return(l[n:] + l[:n])

def prepare_paragraph(doc):
    p = doc.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.left_indent = Cm(1.5)
    paragraph_format.first_line_indent = Cm(-1.5)
    paragraph_format.space_after = Pt(0)
    paragraph_format.keep_with_next = True
    tab_stops = paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(Cm(1.5))
    return(p)


def write_markdown_paragraph(p, text):
    q_wordlist = text.split()

    use_italics = False
    use_bold = False
    for k, word in enumerate(q_wordlist):

        if len(word) > 4 and word[0] == word[1] == word[-1] == word[-2] == "*":  # enclosed bold markdown
            word = word.replace("*","")
            p.add_run(word).bold = True
        elif len(word) > 2 and word[0] == word[-1] == "*":  # enclosed italics markdown
            word = word.replace("*","")
            p.add_run(word).italic = True
        elif len(word) > 2 and word[0] == word[1] == "*":  # start bold markdown
            use_bold = True
            word = word.replace("*","")
            p.add_run(word).bold = True
        elif len(word) > 2 and word[-1] == word[-2] == "*":  # end bold markdown
            use_bold = False
            word = word.replace("*","")
            p.add_run(word).bold = True
        elif len(word) > 1 and word[0] == "*":  # start italics markdown
            use_italics = True
            word = word.replace("*","")
            p.add_run(word).italic = True
        elif len(word) > 1 and word[-1] == "*":  # end italics markdown
            use_italics = False
            word = word.replace("*","")
            p.add_run(word).italic = True
        else:
            if use_bold:
                p.add_run(word).bold = True
            elif use_italics:
                p.add_run(word).italic = True
            else:
                p.add_run(word)
        if k != len(q_wordlist) - 1:
            p.add_run(" ")




def write_answer_file(answer_list):
    # Create a workbook and add a worksheet.
    workbook = xlsxwriter.Workbook('answer_sheet.xlsx')
    worksheet = workbook.add_worksheet()

    # Some data we want to write to the worksheet.
    expenses = (
        ['Rent', 1000],
        ['Gas', 100],
        ['Food', 300],
        ['Gym', 50],
    )

    # Start from the first cell. Rows and columns are zero indexed.
    row = 0
    col = 0

    # Iterate over the data and write it out row by row.
    # for item, cost in (answer_list):
    #     worksheet.write(row, col, item)
    #     worksheet.write(row, col + 1, cost)
    #     row += 1

    # Iterate over the data and write it out row by row.
    for v1, v2, ans in (answer_list):
        worksheet.write(row, col, v1, v2, ans)
        row += 1

    workbook.close()
